import streamlit as st
from pdf_utils import pdf_extract, format_cover_letter
from docx import Document
from fpdf import FPDF
import tempfile
import unicodedata

def clean_text(text):
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')

# App config
st.set_page_config(page_title="Cover Letter Generator", layout="wide")
st.title("📄 AI-Powered Cover Letter Generator")

st.markdown("Upload your **resume**, enter the **job description**, and generate a cover letter using Gemini.")

# Sidebar - User Input
with st.sidebar:
    st.header("Your Details")
    name = st.text_input("Your Name")
    address11 = st.text_input("Address Line 1")
    address12 = st.text_input("Address Line 2")
    address13 = st.text_input("Address Line 3")

    st.header("Recruiter Details")
    name2 = st.text_input("Recruiter Name (e.g. Hiring Manager)")
    designation = st.text_input("Designation (e.g. HR Manager)")
    company = st.text_input("Company Name")
    address21 = st.text_input("Company Address Line 1")
    address22 = st.text_input("Company Address Line 2")
    address23 = st.text_input("Company Address Line 3")

    n_paras = st.number_input("Number of Paragraphs", min_value=1, max_value=6, value=3, step=1)

# Main Area
col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("Upload Resume (PDF)", type=["pdf"])

with col2:
    job_description = st.text_area("Paste Job Description Here", height=300)

# Cover letter preview session state
if "cover_letter" not in st.session_state:
    st.session_state.cover_letter = ""

# Generate button
if st.button("🚀 Generate Cover Letter"):
    if resume_file and job_description:
        with st.spinner("Generating your cover letter..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(resume_file.read())
                resume_text = pdf_extract(tmp_file.name)

            if not resume_text:
                st.error("Could not extract text from the resume. Please try another PDF.")
            else:
                # Generate and format
                st.session_state.cover_letter = format_cover_letter(
                    resume=resume_text,
                    job_description=job_description,
                    name=name,
                    address11=address11.strip(),
                    address12=address12.strip(),
                    address13=address13.strip(),
                    name2=name2,
                    designation=designation,
                    company=company,
                    address21=address21.strip(),
                    address22=address22.strip(),
                    address23=address23.strip(),
                    n_paras=n_paras
                )
                st.success("✅ Cover letter generated!")

    else:
        st.warning("Please upload resume and enter job description.")

# Editable Preview
st.subheader("📝 Cover Letter Preview (Editable)")
edited_letter = st.text_area("Make any changes you want below before exporting:",
                             value=st.session_state.cover_letter,
                             height=400, key="edited_letter")

# Regenerate button
if st.button("🔁 Regenerate"):
    if resume_file and job_description:
        with st.spinner("Regenerating..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(resume_file.read())
                resume_text = pdf_extract(tmp_file.name)

            st.session_state.cover_letter = format_cover_letter(
                resume=resume_text,
                job_description=job_description,
                name=name,
                address11=address11.strip(),
                address12=address12.strip(),
                address13=address13.strip(),
                name2=name2,
                designation=designation,
                company=company,
                address21=address21.strip(),
                address22=address22.strip(),
                address23=address23.strip(),
                n_paras=n_paras
            )
            st.success("🔁 Content regenerated!")
    else:
        st.warning("Upload resume and job description before regenerating.")

# Export buttons
st.subheader("📤 Export")
export_format = st.selectbox("Choose Format", ["Word (.docx)", "PDF (.pdf)"])
if st.button("💾 Export Cover Letter"):
    final_letter = edited_letter.strip()

    if export_format == "Word (.docx)":
        doc = Document()
        for line in final_letter.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph()  # empty line
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            st.download_button("📥 Download Word File", data=open(tmp.name, "rb"), file_name="cover_letter.docx")



    elif export_format == "PDF (.pdf)":
        pdf = FPDF(format='A4')
        pdf.set_auto_page_break(auto=True, margin=25.4)
        pdf.set_left_margin(25.4)
        pdf.set_right_margin(25.4)
        pdf.set_top_margin(25.4)
        pdf.add_page()
        pdf.set_font("Times", size=11)
        line_height = 5.5
        cleaned_text = clean_text(final_letter)
        for line in cleaned_text.split("\n"):
            pdf.multi_cell(0, line_height, line.strip())
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            pdf.output(tmp.name)
            st.download_button("📥 Download PDF File", data=open(tmp.name, "rb"), file_name="cover_letter.pdf")